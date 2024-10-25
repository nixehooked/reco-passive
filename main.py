import os
import pyperclip
import dns.resolver
import dns.exception
import socket
import shodan
import requests
from googletrans import Translator
from docx import Document
import whois
target = input(
    "Veuillez choisir une cible pour énumérer les sous domaines (e.g www.exemple.com): ")
file = input("Donnez un nom au fichier: ")

# Ouvre le modèle Word
doc = Document('template.docx')

vuln_critical = 0
vuln_high = 0
vuln_medium = 0
vuln_low = 0
vuln_major = 0
vuln_minor = 0
vuln_total = 0

banner_subscraper = '''

    ---------------------------------------------------------------------------------
                            ENUMERATION DES SOUS DOMAINES
    ---------------------------------------------------------------------------------

    '''

banner_shodan = '''

    ---------------------------------------------------------------------------------
                                ENUMERATION SHODAN
    ---------------------------------------------------------------------------------

                    '''

# Clé API Shodan
SHODAN_API_KEY = 'cdngF34tKPzqZMe1zo5scYPCyPShPvwT'

# Clé API NVD
NVD_API_KEY = '4bcc7d06-e7e4-4870-9b05-bcdf517e771a'


# Enumération de directories
def dirb_enum():
    os.system(f'dirb {target}')


# Enumération des sous domaines
def subdomains_enum():
    """
    Execute subscraper et récupére chaque targets dans un .txt
    """
    os.system(f'python subscraper.py {target} -r {file}.txt')

    with open(f'{file}.txt', 'r') as f:
        lines = f.readlines()
        res = []
        for sub in lines:
            res.append(sub.replace("\n", ""))

    return res


# Traduction
def translate_description(description):
    """ Utilise l'API google translate pour traduire les descriptions des CVE de l'API NVD

    Args:
        description: Description retourné par l'API NVD

    Returns:
        _type_: Description traduite
    """
    try:
        translator = Translator(service_urls=['translate.google.com'])
        translation = translator.translate(description, dest='fr')
        if translation.text:
            return translation.text
        else:
            # Si la traduction est vide, renvoyer la description d'origine
            return description
    except Exception as e:
        print("Erreur lors de la traduction :", str(e))
        return description


def is_cname(domain):
    """
    Vérifie si le domaine est un CNAME en utilisant la bibliothèque dns.resolver de Python.
    """
    try:
        answers = dns.resolver.resolve(domain, 'CNAME')
        return True
    except dns.exception.DNSException:
        return False


def filter_subdomains(subdomains):
    """
    Filtre une liste de sous-domaines pour enlever tous les CNAME.
    """
    return [subdomain for subdomain in subdomains if not is_cname(subdomain)]


# Fonction pour récupérer les détails d'une CVE à partir de l'API NVD
def get_cve_details(cve_id):
    """_summary_

    Args:
        cve_id (_type_): Nom de la CVE

    Returns:
        _type_: Description traduite et le score CVSS
    """
    url = f'https://services.nvd.nist.gov/rest/json/cve/1.0/{cve_id}'
    headers = {'api_key': NVD_API_KEY}
    response = requests.get(url, headers=headers)
    if response.status_code == 200:
        data = response.json()
        cve_items = data.get('result', {}).get('CVE_Items', [])
        if cve_items:
            cve_details = cve_items[0]['cve']['description']['description_data'][0]['value']
            translated_description = translate_description(cve_details)
            # Récupérer le score CVSSv3 s'il est disponible
            cvss_v3 = cve_items[0]['impact'].get('baseMetricV3')

            # Récupérer le score CVSSv2 s'il existe et CVSSv3 n'est pas disponible
            if cvss_v3 is None:
                cvss_v2 = cve_items[0]['impact'].get('baseMetricV2')
                if cvss_v2 is not None:
                    cvss_score = cvss_v2.get('cvssV2', {}).get('baseScore')

            # Si le score CVSSv3 est disponible, récupérer le score CVSSv3
            if cvss_v3 is not None:
                cvss_score = cvss_v3.get('cvssV3', {}).get('baseScore')
            return translated_description, cvss_score
        else:
            return 'Détails indisponibles', None, 'null'
    else:
        return 'Détails indisponibles', None


# Fonction pour obtenir le niveau de criticité en fonction du score CVSS
def get_criticite(score_cvss):
    """
    Retourne un niveau de criticité selon le score CVSS
    """
    if score_cvss >= 9.0:
        return "Critique"
    elif score_cvss >= 7.0:
        return "Élevée"
    elif score_cvss >= 4.0:
        return "Moyenne"
    else:
        return "Faible"


r = subdomains_enum()
filtered_subdomains = filter_subdomains(r)
# Trouver le paragraphe contenant le repère
for paragraph in doc.paragraphs:
    if '{SUBDOMAINS}' in paragraph.text:
        # Remplacer le repère par les sous-domaines
        paragraph.text = paragraph.text.replace(
            '{SUBDOMAINS}', '\n'.join(filtered_subdomains))

with open(f'{file}.txt', 'r') as f:
    lines = f.readlines()
    res = []
    for sub in lines:
        res.append(sub.replace("\n", ""))
    targets = ', '.join(res)
    pyperclip.copy(targets)

adresses_ip = set()
sous_domaines = []

for sous_domaine in filtered_subdomains:
    try:
        adresse_ip = socket.gethostbyname(sous_domaine)
        adresses_ip.add(adresse_ip)
        sous_domaines.append(sous_domaine)
        print(sous_domaine + ' : ' + adresse_ip)
    except socket.gaierror:
        print(sous_domaine + ' : Impossible de résoudre le nom de domaine')

# Connexion à l'API Shodan
api = shodan.Shodan(SHODAN_API_KEY)

with open(f"{file}.txt", 'a') as f:
    f.write(banner_shodan)

for ip, sous_domaines in zip(adresses_ip, sous_domaines):
    try:
        w = whois.whois(ip)
        print(f"Registrar: {w.registrar}")
        print(f"Organisation: {w.org}")
        print(f"Pays: {w.country}")
        print(f"Ville: {w.city}")
        print(f"Nom du propriétaire: {w.name}")
        print(f"Téléphone: {w.phone}")
        # Recherche des informations sur l'adresse IP
        resultats = api.host(ip)
        # Affichage des résultats
        print('Résultats de recherche pour', ip, ':')
        print('Domaine:', resultats['domains'])
        print('Hostnames:', resultats['hostnames'])
        print('IP:', resultats['ip_str'])
        data = resultats['data']
        for service in data:
            print('Port:', service['port'])

            if 'vulns' in service:
                vulns = service['vulns']
                print('Vulnérabilités:')
                for vuln in vulns:
                    print('- CVE:', vuln)
                    cve_details, cvss_score = get_cve_details(vuln)
                    print('Description:', cve_details)
                    print('CVSS Score:', cvss_score, '\n')
            else:
                print('Aucune vulnérabilité trouvée')

        print('---')

        new_paragraphs = []
        for paragraph in doc.paragraphs:
            if '{CVE_CONTENT}' not in paragraph.text:
                new_paragraphs.append(paragraph)

        # Créer un nouveau document
        new_doc = Document()

        # Copier les paragraphes existants dans le nouveau document
        for paragraph in new_paragraphs:
            new_doc.add_paragraph(paragraph.text, style=paragraph.style)

        # Itérer sur les informations des CVE

        # Dupliquer le contenu des CVE avec les en-têtes
        doc.add_paragraph(
            text=f'Résultats de recherche pour {sous_domaines} ({ip})', style='Heading 2')

        doc.add_paragraph("Registrar:", style='Style1')
        if w.registrar:
            doc.add_paragraph(w.registrar, style='List Paragraph')
        else:
            doc.add_paragraph("Non disponible", style='List Paragraph')

        doc.add_paragraph("Organisation:", style='Style1')
        if w.org:
            doc.add_paragraph(w.org, style='List Paragraph')
        else:
            doc.add_paragraph("Non disponible", style='List Paragraph')

        doc.add_paragraph("Pays:", style='Style1')
        if w.country:
            doc.add_paragraph(w.country, style='List Paragraph')
        else:
            doc.add_paragraph("Non disponible", style='List Paragraph')

        doc.add_paragraph("Ville:", style='Style1')
        if w.city:
            doc.add_paragraph(w.city, style='List Paragraph')
        else:
            doc.add_paragraph("Non disponible", style='List Paragraph')

        doc.add_paragraph("Nom du propriétaire:", style='Style1')
        if w.name:
            doc.add_paragraph({w.name}, style='List Paragraph')
        else:
            doc.add_paragraph("Non disponible", style='List Paragraph')

        doc.add_paragraph("Email:", style='Style1')
        if w.emails:
            for email in w.emails:
                doc.add_paragraph(f"- {email}", style='List Paragraph')
        else:
            doc.add_paragraph("Non disponible", style='List Paragraph')

        doc.add_paragraph("Téléphone:", style='Style1')
        if w.phone:
            doc.add_paragraph({w.phone}, style='List Paragraph')
        else:
            doc.add_paragraph("Non disponible", style='List Paragraph')
        doc.add_paragraph('Domaines:', style='Style1')
        for domain in resultats['domains']:
            doc.add_paragraph(f"- {domain}")
        doc.add_paragraph('Hostnames:', style='Style1')
        for host in resultats['hostnames']:
            doc.add_paragraph(f"- {host}")
        doc.add_paragraph('Ports:', style='Normal Ports')
        for service in data:
            if 'name' in service:
                service_name = service['name']
            elif 'product' in service:
                service_name = service['product']
            else:
                service_name = "Non reconnu"
            doc.add_paragraph(text=f"- {service['port']} | {service_name}", style='Heading 3 Port')
            if 'vulns' in service:
                vulns = service['vulns']
                doc.add_paragraph('Vulnérabilités:')
                for vuln in vulns:
                    doc.add_paragraph(
                        text=f'- CVE: {vuln}\n', style='Heading 4 CVE')
                    cve_details, cvss_score = get_cve_details(vuln)
                    doc.add_paragraph(text=f'Description: {cve_details}')
                    doc.add_paragraph(text=f'Score CVSS: {cvss_score}')
                    if cvss_score is not None:
                        cvss_score_float = float(cvss_score)
                    else:
                        cvss_score_float = 0.0
                    criticite = get_criticite(cvss_score_float)
                    if criticite == 'Critique' or criticite == 'Élevée':
                        if criticite == 'Critique':
                            vuln_critical += 1
                            doc.add_paragraph(
                                text=f"Criticite: {criticite}", style='Vulnérabilité 04 Critique')
                        elif criticite == 'Élevée':
                            vuln_high += 1
                            doc.add_paragraph(
                                text=f"Criticite: {criticite}", style='Vulnérabilité 03 élevée')
                        vuln_major += 1
                    elif criticite == 'Moyenne' or criticite == 'Faible':
                        if criticite == 'Moyenne':
                            vuln_medium += 1
                            doc.add_paragraph(
                                text=f"Criticite: {criticite}", style='Vulnérabilité 02 Moyenne')
                        elif criticite == 'Faible':
                            vuln_low += 1
                            doc.add_paragraph(
                                text=f"Criticite: {criticite}", style='Vulnérabilité 01 faible')
                        vuln_minor += 1
            else:
                doc.add_paragraph(text='Aucune vulnérabilité trouvée')
            vuln_total = vuln_major + vuln_minor

            # Remplacer la balise {CVE_CONTENT} par le contenu des CVE
            for paragraph in doc.paragraphs:
                if '{CVE_CONTENT}' in paragraph.text:
                    # Supprimer la balise {CVE_CONTENT}
                    paragraph.text = paragraph.text.replace('{CVE_CONTENT}', '')
                    # Insérer le contenu des CVE ici en utilisant doc.add_paragraph() selon vos besoins

        for paragraph in doc.paragraphs:
            if '{TARGET}' in paragraph.text:
                paragraph.text = paragraph.text.replace('{TARGET}', target)

    except shodan.APIError as e:
        print('Erreur lors de la recherche de', adresse_ip, ':', e)

for paragraph in doc.paragraphs:
    critical = str(vuln_critical)
    high = str(vuln_high)
    medium = str(vuln_medium)
    low = str(vuln_low)
    if '{NOTE}' in paragraph.text:
        if vuln_major == 0 and vuln_minor <= 4:
            paragraph.text = paragraph.text.replace('{NOTE}', "A+")
        elif vuln_major <= 2 and vuln_minor <= 8 and vuln_total <= 10:
            paragraph.text = paragraph.text.replace('{NOTE}', "A")
        elif vuln_major <= 4 and vuln_minor <= 12 and vuln_total <= 15:
            paragraph.text = paragraph.text.replace('{NOTE}', "B")
        elif vuln_major <= 6 and vuln_minor <= 16 and vuln_total <= 20:
            paragraph.text = paragraph.text.replace('{NOTE}', "C")
        else:
            paragraph.text = paragraph.text.replace('{NOTE}', "D")
    if '{VULN_CRITICAL}' in paragraph.text:
        paragraph.text = paragraph.text.replace('{VULN_CRITICAL}', critical)
    if '{VULN_HIGH}' in paragraph.text:
        paragraph.text = paragraph.text.replace('{VULN_HIGH}', high)
    if '{VULN_MEDIUM}' in paragraph.text:
        paragraph.text = paragraph.text.replace('{VULN_MEDIUM}', medium)
    if '{VULN_LOW}' in paragraph.text:
        paragraph.text = paragraph.text.replace('{VULN_LOW}', low)

# Enregistrer le document Word modifié
doc.save(f'{file}_report.docx')